"""
CHECKPOINT 4 – Template Exporter Test
======================================

Verifies that TemplateExporter correctly produces filled Word and Excel
files from a mock TemplateProfile and a list of generated timetable slots.

Run directly (no DB needed):
    cd backend
    python tests/run_exporter_tests.py
"""

import io
import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from app.utils.template_exporter import TemplateExporter

PASS = "[PASS]"
FAIL = "[FAIL]"
results = {"passed": 0, "failed": 0}


def check(name: str, condition: bool, detail: str = ""):
    if condition:
        print(f"  {PASS} {name}")
        results["passed"] += 1
    else:
        print(f"  {FAIL} {name}" + (f"\n       detail: {detail}" if detail else ""))
        results["failed"] += 1


# ---------------------------------------------------------------------------
# Mock data
# ---------------------------------------------------------------------------

class FakeProfile:
    file_type = "xlsx"
    containers = [
        {"session_type": "lecture",   "day": "Monday",  "start_hour": 8,  "end_hour": 10, "group_label": "AEN-3"},
        {"session_type": "practical", "day": "Tuesday", "start_hour": 14, "end_hour": 17, "group_label": "GEN-2"},
        {"session_type": "tutorial",  "day": "Wednesday", "start_hour": 10, "end_hour": 12, "group_label": "EEE-5"},
    ]


class FakeDocxProfile:
    file_type = "docx"
    containers = FakeProfile.containers


MOCK_SLOTS = [
    {
        "id": 1,
        "day_of_week": 0,           # Monday
        "start_time": "08:00:00",
        "end_time": "10:00:00",
        "session_type": "lecture",
        "course_code": "CEE 3101",
        "lecturer_name": "Dr. Banda",
        "room_name": "NLR1",
    },
    {
        "id": 2,
        "day_of_week": 1,           # Tuesday
        "start_time": "14:00:00",
        "end_time": "17:00:00",
        "session_type": "practical",
        "course_code": "GEE 2201",
        "lecturer_name": "Eng. Mwale",
        "room_name": "Lab-A",
    },
]


# ---------------------------------------------------------------------------
# Excel export tests
# ---------------------------------------------------------------------------

def test_excel_export():
    print("\n=== Excel Export Tests ===")
    profile = FakeProfile()
    exporter = TemplateExporter(profile, MOCK_SLOTS)

    try:
        data, mime = exporter.export_bytes()
    except Exception as e:
        import traceback
        print(f"  {FAIL} export_bytes() raised: {e}")
        traceback.print_exc()
        results["failed"] += 5
        return

    check("export_bytes returns bytes",         isinstance(data, bytes))
    check("data is non-empty",                  len(data) > 1000, f"size={len(data)}")
    check("mime is xlsx type",                  "spreadsheet" in mime)

    # Verify the file is a valid xlsx by loading it
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data))
        ws = wb.active
        check("workbook loads successfully",   True)
        check("has more than 1 row",           ws.max_row > 1, f"rows={ws.max_row}")
        check("has TIME header in A1",         ws['A1'].value == "TIME", f"A1={ws['A1'].value}")
        check("has Monday header",             any(ws.cell(1, c).value == "Monday" for c in range(1,8)))
    except Exception as e:
        check("workbook loads successfully",   False, str(e))


# ---------------------------------------------------------------------------
# Word export tests
# ---------------------------------------------------------------------------

def test_docx_export():
    print("\n=== Word Export Tests ===")
    profile = FakeDocxProfile()
    exporter = TemplateExporter(profile, MOCK_SLOTS)

    try:
        data, mime = exporter.export_bytes()
    except Exception as e:
        import traceback
        print(f"  {FAIL} export_bytes() raised: {e}")
        traceback.print_exc()
        results["failed"] += 5
        return

    check("export_bytes returns bytes",        isinstance(data, bytes))
    check("data is non-empty",                 len(data) > 1000, f"size={len(data)}")
    check("mime is docx type",                 "wordprocessingml" in mime)

    # Verify the file is a valid docx by loading it
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        check("Document loads successfully",   True)
        check("Has at least 1 table",          len(doc.tables) >= 1, f"tables={len(doc.tables)}")
        table = doc.tables[0]
        check("Table has TIME in first cell",  table.rows[0].cells[0].text == "TIME",
              f"got: {table.rows[0].cells[0].text}")
    except Exception as e:
        check("Document loads successfully",   False, str(e))


# ---------------------------------------------------------------------------
# Slot content is painted into cells
# ---------------------------------------------------------------------------

def test_slot_content_in_excel():
    print("\n=== Cell Content Painting Test ===")
    profile = FakeProfile()
    exporter = TemplateExporter(profile, MOCK_SLOTS)
    data, _ = exporter.export_bytes()

    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data))
    ws = wb.active

    # Scan all cells for the course code
    found_lecture  = any("CEE 3101" in str(cell.value or "") for row in ws for cell in row)
    found_practical = any("GEE 2201" in str(cell.value or "") for row in ws for cell in row)
    check("CEE 3101 (lecture) appears in Excel", found_lecture,
          "Slot was not painted into the Excel file")
    check("GEE 2201 (practical) appears in Excel", found_practical,
          "Slot was not painted into the Excel file")


# ---------------------------------------------------------------------------
# Extra lookup (slot objects without string fields)
# ---------------------------------------------------------------------------

def test_extra_lookup():
    print("\n=== Extra Lookup Test ===")
    # Slots with IDs only; extra_lookup provides the strings
    bare_slots = [{"id": 10, "day_of_week": 0, "start_time": "08:00:00",
                   "end_time": "10:00:00", "session_type": "lecture"}]
    extra = {10: {"course_code": "CEE-EXTRA", "lecturer_name": "Prof. Test", "room_name": "LH1"}}

    profile = FakeProfile()
    exporter = TemplateExporter(profile, bare_slots, extra_lookup=extra)
    data, _ = exporter.export_bytes()

    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data))
    ws = wb.active
    found = any("CEE-EXTRA" in str(cell.value or "") for row in ws for cell in row)
    check("Extra lookup name appears in output", found)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    test_excel_export()
    test_docx_export()
    test_slot_content_in_excel()
    test_extra_lookup()

    total = results["passed"] + results["failed"]
    bar = "=" * 55
    print(f"\n{bar}")
    print(f"CHECKPOINT 4 RESULTS: {results['passed']}/{total} passed")
    if results["failed"] > 0:
        print(f"  {results['failed']} FAILURES - review output above")
        sys.exit(1)
    else:
        print("  ALL TESTS PASSED")
        sys.exit(0)
