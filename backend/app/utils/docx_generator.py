from typing import Any, Dict, List

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DAYS: List[str] = ["MONDAY", "TUESDAY", "WEDNESDAY", "THURSDAY", "FRIDAY"]
HOURS: List[str] = [f"{h:02d}:00" for h in range(7, 19)]
COLUMN_KEYS: List[str] = [
    "GEN LG1",
    "GEN LG2",
    "3-AEN",
    "3-CEE",
    "3-EEE",
    "3-GEE",
    "3-MEC",
    "4-AEN",
    "4-CEE",
    "4-EEE",
    "4-GEE",
    "4-MEC",
    "5-AEN",
    "5-CEE",
    "5-EEE",
    "5-GEE",
    "5-MEC",
]
DISPLAY_HEADERS: List[str] = [
    "GEN LG1",
    "GEN LG2",
    "AEN",
    "CEE",
    "EEE",
    "GEE",
    "MEC",
    "AEN",
    "CEE",
    "EEE",
    "GEE",
    "MEC",
    "AEN",
    "CEE",
    "EEE",
    "GEE",
    "MEC",
]

COLOR_DARK_BLUE = "0B3A70"
COLOR_ORANGE = "E98A15"
COLOR_LIGHT_BLUE = "EAF2FB"
COLOR_LIGHT_GRAY = "F6F8FA"

ROOM_KEY_NOTES = [
    "MLT - School of Mines Lecture Theatre",
    "LT - Lecture Theatre, School of Engineering",
    "DR - Drawing Room",
    "CR - Computer Lab",
    "CELAB - Civil Engineering Laboratory",
    "EELAB - Electrical Engineering Laboratory",
]


class DocxGenerator:
    def __init__(self, output_path: str):
        self.output_path = output_path
        self.document = Document()
        self._setup_document()

    def _setup_document(self) -> None:
        section = self.document.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(42.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(0.8)
        section.right_margin = Cm(0.8)

        style = self.document.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(7.5)

    def generate(self, data: Dict[str, Any]) -> str:
        self._add_header(data)
        grid = data.get("grid_data", {})

        for index, day in enumerate(DAYS):
            self._add_day_section(day, grid.get(day, {}))
            if index < len(DAYS) - 1:
                self.document.add_page_break()

        self.document.add_page_break()
        self._add_footer_keys()
        self.document.save(self.output_path)
        return self.output_path

    def _add_header(self, data: Dict[str, Any]) -> None:
        university_name = data.get("university_name", "UNIVERSITY")
        timetable_name = data.get("timetable_name", "TIMETABLE")
        semester = data.get("semester", "")
        year = data.get("year", "")
        half_text = "First Half" if data.get("academic_half") == "first_half" else "Second Half"

        title = self.document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(university_name.upper())
        run.bold = True
        run.font.size = Pt(16)

        subtitle = self.document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(timetable_name)
        run.bold = True
        run.font.size = Pt(12)

        meta = self.document.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta.add_run(f"{semester} • {half_text} • {year} Academic Year")
        run.font.size = Pt(10)
        run.italic = True

        spacer = self.document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(4)

    def _add_day_section(self, day: str, day_data: Dict[str, Any]) -> None:
        day_heading = self.document.add_paragraph()
        day_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = day_heading.add_run(day)
        run.bold = True
        run.font.size = Pt(11)

        table = self.document.add_table(rows=2, cols=18)
        table.style = "Table Grid"
        table.autofit = False

        widths = [Cm(1.9)] + [Cm(2.15)] * 17
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        self._populate_header_rows(table)

        for hour in HOURS:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{hour}-{int(hour[:2]) + 1:02d}:00"
            self._style_cell(row_cells[0], COLOR_LIGHT_BLUE, bold=True, font_size=7.5)

            hour_data = day_data.get(hour, {})
            for idx, key in enumerate(COLUMN_KEYS, start=1):
                slots = hour_data.get(key, [])
                if slots:
                    slot = slots[0]
                    row_cells[idx].text = (
                        f"{slot.get('course_code', '')}\n"
                        f"{slot.get('room_name', '')}\n"
                        f"{slot.get('lecturer_name', '')}"
                    ).strip()
                else:
                    row_cells[idx].text = ""
                shaded = COLOR_LIGHT_GRAY if len(table.rows) % 2 == 0 else "FFFFFF"
                self._style_cell(row_cells[idx], shaded, font_size=7)

        self.document.add_paragraph()

    def _populate_header_rows(self, table) -> None:
        row0 = table.rows[0].cells
        row1 = table.rows[1].cells

        row0[0].text = "HOURS"
        row0[1].text = "2ND YEAR"
        row0[3].text = "3RD YEAR"
        row0[8].text = "4TH YEAR"
        row0[13].text = "5TH YEAR"

        row0[1].merge(row0[2])
        row0[3].merge(row0[7])
        row0[8].merge(row0[12])
        row0[13].merge(row0[17])

        for idx in [0, 1, 3, 8, 13]:
            self._style_cell(row0[idx], COLOR_DARK_BLUE, bold=True, font_size=8.5, color="FFFFFF")

        headers = [""] + DISPLAY_HEADERS
        for idx, label in enumerate(headers):
            row1[idx].text = label
            self._style_cell(row1[idx], COLOR_ORANGE, bold=True, font_size=8, color="FFFFFF")

    def _add_footer_keys(self) -> None:
        heading = self.document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = heading.add_run("Room Key")
        run.bold = True
        run.font.size = Pt(12)

        for note in ROOM_KEY_NOTES:
            paragraph = self.document.add_paragraph(style="List Bullet")
            paragraph.add_run(note)

    def _style_cell(
        self,
        cell,
        fill: str,
        *,
        bold: bool = False,
        font_size: float = 7.5,
        color: str = "000000",
    ) -> None:
        self._set_cell_background(cell, fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.bold = bold
                run.font.size = Pt(font_size)
                run.font.color.rgb = self._hex_to_rgb(color)

    def _set_cell_background(self, cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    @staticmethod
    def _hex_to_rgb(value: str):
        from docx.shared import RGBColor

        value = value.strip().lstrip("#")
        return RGBColor.from_string(value)
