"""
Template Exporter
=================

Takes a list of generated TimetableSlot results and a TemplateProfile,
then paints the course/lecturer/room information back into the same
cell positions used in the original template file (Word or Excel).

Outputs a new file that looks exactly like the coordinator's original
timetable format, but filled with the generated schedule.

Usage::

    from app.utils.template_exporter import TemplateExporter

    exporter = TemplateExporter(profile, slots, output_path="filled_timetable.xlsx")
    exporter.export()

    # Or get raw bytes (for streaming to frontend download):
    data = exporter.export_bytes()
"""

import io
import logging
from datetime import time
from typing import Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

DAYS = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']

# Background fills for each session type (Excel RGB hex)
SESSION_FILL_COLORS = {
    'lecture':   'DDEEFF',  # light blue
    'practical': 'DDFFDD',  # light green
    'tutorial':  'FFF3CD',  # light amber
}

# Text displayed in each filled cell
DEFAULT_CELL_TEMPLATE = "{course_code}\n{lecturer}\n{room}"


# ---------------------------------------------------------------------------
# Data class for a generated slot (flexible dict or object)
# ---------------------------------------------------------------------------

def _get(slot, key: str, default=None):
    """Works with both dict slots and ORM objects."""
    if isinstance(slot, dict):
        return slot.get(key, default)
    return getattr(slot, key, default)


# ---------------------------------------------------------------------------
# Main exporter
# ---------------------------------------------------------------------------

class TemplateExporter:
    """
    Export generated timetable slots back into an original template file
    (Excel or Word), using the coordinate map from a TemplateProfile.

    Parameters
    ----------
    profile : TemplateProfile ORM model instance (or dict-like)
    slots   : list of TimetableSlot ORM objects or dicts, each with:
              day_of_week (int 0-4), start_time (time), end_time (time),
              session_type (str), course_code (str), lecturer_name (str),
              room_name (str).
    """

    def __init__(self, profile, slots: List, extra_lookup: Optional[Dict] = None):
        """
        Parameters
        ----------
        profile      : TemplateProfile ORM object with .containers, .file_type
        slots        : generated timetable slots (list of dicts or ORM objects)
        extra_lookup : optional dict {slot_id: {course_code, lecturer_name, room_name}}
                       useful when slot objects don't carry full string names.
        """
        self.profile = profile
        self.slots = slots
        self.extra_lookup = extra_lookup or {}

        # Build an index of containers for fast lookup
        # Key: (day_idx, start_hour, session_type) → container  (first match)
        self._container_index: Dict[Tuple[int, int, str], dict] = {}
        self._build_container_index()

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def export_bytes(self) -> Tuple[bytes, str]:
        """
        Return (file_bytes, mime_type) for the filled timetable.
        mime_type is suitable for an HTTP Content-Type header.
        """
        file_type = getattr(self.profile, 'file_type', 'xlsx')

        if file_type == 'docx':
            data = self._export_docx()
            mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            data = self._export_excel()
            mime = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'

        return data, mime

    def export(self, output_path: str) -> None:
        """Write the filled file to disk at *output_path*."""
        data, _ = self.export_bytes()
        with open(output_path, 'wb') as f:
            f.write(data)
        logger.info("Exported filled timetable to %s", output_path)

    # ------------------------------------------------------------------
    # Index building
    # ------------------------------------------------------------------

    def _build_container_index(self) -> None:
        """Map (day_idx, start_hour, session_type) → container dict."""
        containers = getattr(self.profile, 'containers', None) or []
        for c in containers:
            day_name     = c.get('day', '')
            start_hour   = c.get('start_hour')
            session_type = c.get('session_type')

            if day_name not in DAYS or start_hour is None or session_type is None:
                continue

            day_idx = DAYS.index(day_name)
            key = (day_idx, start_hour, session_type)
            if key not in self._container_index:
                self._container_index[key] = c

    # ------------------------------------------------------------------
    # Slot → cell-label helpers
    # ------------------------------------------------------------------

    def _slot_label(self, slot) -> str:
        """Build the string to write into a timetable cell."""
        slot_id = _get(slot, 'id')
        extra   = self.extra_lookup.get(slot_id, {})

        course_code   = extra.get('course_code')   or _get(slot, 'course_code',   'TBD')
        lecturer_name = extra.get('lecturer_name') or _get(slot, 'lecturer_name', 'TBD')
        room_name     = extra.get('room_name')      or _get(slot, 'room_name',     'TBD')

        return DEFAULT_CELL_TEMPLATE.format(
            course_code=course_code,
            lecturer=lecturer_name,
            room=room_name,
        )

    def _slot_lookup(self) -> Dict[Tuple[int, int, str], List]:
        """
        Build a mapping of (day_idx, start_hour, session_type) → list of slots.
        Used by exporters to paint slots into matching containers.
        """
        lookup: Dict[Tuple[int, int, str], List] = {}
        for slot in self.slots:
            day_idx      = _get(slot, 'day_of_week', 0)
            start_time   = _get(slot, 'start_time')
            session_type = _get(slot, 'session_type', 'lecture')

            if start_time is None:
                continue

            if isinstance(start_time, time):
                start_hour = start_time.hour
            else:
                # Handle "HH:MM:SS" string
                start_hour = int(str(start_time).split(':')[0])

            key = (day_idx, start_hour, session_type)
            lookup.setdefault(key, []).append(slot)

        return lookup

    # ------------------------------------------------------------------
    # Excel exporter
    # ------------------------------------------------------------------

    def _export_excel(self) -> bytes:
        """
        Build a minimal Excel workbook representing the filled timetable.

        Strategy:
        - Row 0: Day headers
        - Subsequent rows: one per hour (07:00 – 18:00)
        - Cells are coloured by session type and labelled with slot content.

        If a more faithful clone of the original template is needed, store
        the original file bytes in TemplateProfile and use openpyxl's
        load_workbook() instead.
        """
        try:
            import openpyxl
            from openpyxl.styles import PatternFill, Alignment, Font, Border, Side
        except ImportError:
            raise ImportError("openpyxl required: pip install openpyxl")

        slot_lookup = self._slot_lookup()
        containers  = getattr(self.profile, 'containers', []) or []

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Timetable"

        # Header row
        ws.cell(row=1, column=1, value="TIME")
        for col_idx, day in enumerate(DAYS, start=2):
            ws.cell(row=1, column=col_idx, value=day)

        # Time rows (07:00 – 18:00)
        hours = list(range(7, 19))
        for row_i, hour in enumerate(hours, start=2):
            ws.cell(row=row_i, column=1, value=f"{hour:02d}:00")

        # Paint containers
        for c in containers:
            day_name     = c.get('day', '')
            start_hour   = c.get('start_hour')
            end_hour     = c.get('end_hour', start_hour + 1 if start_hour else None)
            session_type = c.get('session_type', 'lecture')

            if day_name not in DAYS or start_hour is None:
                continue

            day_col  = DAYS.index(day_name) + 2   # column (1-indexed, offset by TIME col)
            start_row = (start_hour - 7) + 2       # row offset: hour 7 → row 2
            end_row   = (end_hour - 7) + 2 if end_hour else start_row + 1

            # Find matching slots
            key   = (DAYS.index(day_name), start_hour, session_type)
            slots = slot_lookup.get(key, [])
            label = '\n'.join(self._slot_label(s) for s in slots) if slots else f"[{session_type.upper()}]"

            fill_color = SESSION_FILL_COLORS.get(session_type, 'FFFFFF')
            fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type='solid')

            # Merge cells vertically for multi-hour spans
            if end_row > start_row + 1:
                ws.merge_cells(
                    start_row=start_row, start_column=day_col,
                    end_row=end_row - 1, end_column=day_col
                )

            cell = ws.cell(row=start_row, column=day_col, value=label)
            cell.fill = fill
            cell.alignment = Alignment(wrap_text=True, vertical='top')

        # Auto-size columns roughly
        for col in ws.columns:
            ws.column_dimensions[col[0].column_letter].width = 22

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------
    # Word exporter
    # ------------------------------------------------------------------

    def _export_docx(self) -> bytes:
        """
        Build a Word document with one table representing the filled timetable.

        Columns: TIME | Monday | Tuesday | Wednesday | Thursday | Friday
        Rows:    one per hour 07:00 – 18:00
        Cells:   coloured by session type, labelled with slot content.
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")

        slot_lookup = self._slot_lookup()
        containers  = getattr(self.profile, 'containers', []) or []

        doc = Document()
        doc.add_heading("Generated Timetable", level=1)

        hours = list(range(7, 19))
        n_cols = 1 + len(DAYS)   # TIME + 5 days

        table = doc.add_table(rows=1 + len(hours), cols=n_cols)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "TIME"
        for i, day in enumerate(DAYS):
            header_cells[i + 1].text = day

        # Time column
        for row_i, hour in enumerate(hours, start=1):
            table.rows[row_i].cells[0].text = f"{hour:02d}:00"

        # Paint containers
        for c in containers:
            day_name     = c.get('day', '')
            start_hour   = c.get('start_hour')
            session_type = c.get('session_type', 'lecture')

            if day_name not in DAYS or start_hour is None:
                continue

            day_col  = DAYS.index(day_name) + 1
            row_idx  = (start_hour - 7) + 1

            if row_idx < 1 or row_idx >= len(table.rows):
                continue

            key   = (DAYS.index(day_name), start_hour, session_type)
            slots = slot_lookup.get(key, [])
            label = '\n'.join(self._slot_label(s) for s in slots) if slots else f"[{session_type.upper()}]"

            cell = table.rows[row_idx].cells[day_col]
            cell.text = label

            # Apply background shading
            hex_color = SESSION_FILL_COLORS.get(session_type, 'FFFFFF')
            _set_cell_background(cell, hex_color)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


# ---------------------------------------------------------------------------
# Word cell shading helper
# ---------------------------------------------------------------------------

def _set_cell_background(cell, hex_color: str) -> None:
    """Apply a solid background colour to a python-docx table cell."""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tc_pr.append(shd)
    except Exception:
        pass  # Non-critical: shading is cosmetic
